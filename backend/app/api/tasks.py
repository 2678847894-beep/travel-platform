from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from sqlalchemy import cast, Date
from typing import List, Optional
from datetime import date, datetime, timedelta
from app.core.database import get_db
from app.core.auth import get_current_user, require_admin
from app.models.user import User
from app.models.sop import DailyTask
from app.schemas.schemas import TaskOut, TaskCreate
from pydantic import BaseModel
import io

router = APIRouter(prefix='/api/tasks', tags=['Tasks'])


def _build_task_out(task: DailyTask, query_date: date) -> TaskOut:
    """Compute dynamic fields and return TaskOut."""
    task_date_only = task.task_date.date() if isinstance(task.task_date, datetime) else task.task_date
    completed_date_val = task.completed_date
    is_completed_today = (completed_date_val is not None and completed_date_val == query_date)
    is_overdue = (task_date_only < query_date and not is_completed_today)
    return TaskOut(
        id=task.id,
        title=task.title,
        task_date=task.task_date,
        end_date=task.end_date,
        task_time=task.task_time or "",
        end_time=task.end_time or "",
        location=task.location or "",
        description=task.description or "",
        trip_filter=task.trip_filter or "全部",
        is_completed=is_completed_today,
        is_overdue=is_overdue,
        completed_date=task.completed_date,
        created_at=task.created_at,
    )


@router.get('', response_model=List[TaskOut])
def list_tasks(
    task_date: str = None,
    trip_filter: str = None,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    q = db.query(DailyTask)
    if task_date:
        q = q.filter(cast(DailyTask.task_date, Date) == task_date)
    if trip_filter and trip_filter != 'all':
        q = q.filter(DailyTask.trip_filter == trip_filter)

    tasks = q.order_by(DailyTask.task_time, DailyTask.id).all()
    qt_date = date.fromisoformat(task_date) if task_date else date.today()

    result = [_build_task_out(t, qt_date) for t in tasks]
    # Sort: overdue first, then by task_time + id
    result.sort(key=lambda t: (not t.is_overdue, t.task_time or "99:99", t.id))
    return result


@router.post('', response_model=TaskOut)
def create_task(
    body: TaskCreate,
    db: Session = Depends(get_db),
    _: User = Depends(require_admin),
):
    task = DailyTask(**body.model_dump())
    # Default end_date = task_date + 365 days if not provided
    if not task.end_date:
        td = task.task_date.date() if isinstance(task.task_date, datetime) else task.task_date
        task.end_date = datetime.combine(td + timedelta(days=365), datetime.min.time())
    db.add(task)
    db.commit()
    db.refresh(task)
    qt_date = task.task_date.date() if isinstance(task.task_date, datetime) else task.task_date
    return _build_task_out(task, qt_date)


@router.post('/{task_id}/toggle')
def toggle_task(
    task_id: int,
    toggle_date: str = None,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
):
    task = db.query(DailyTask).filter(DailyTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404)

    td = date.today()
    if toggle_date:
        try:
            td = date.fromisoformat(toggle_date)
        except ValueError:
            pass

    if task.completed_date == td:
        # Already completed today → uncomplete
        task.completed_date = None
        task.is_completed = False
    else:
        # Complete for today
        task.completed_date = td
        task.is_completed = True

    db.commit()
    return {'ok': True}


@router.delete('/{task_id}')
def delete_task(
    task_id: int,
    db: Session = Depends(get_db),
    _: User = Depends(require_admin),
):
    task = db.query(DailyTask).filter(DailyTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404)
    db.delete(task)
    db.commit()
    return {'ok': True}


# ── AI 识别导入预览 ──

_COL_CANDIDATES_TITLE = ['任务', '名称', 'title', '事项', '内容', 'task', 'name', 'item', 'event']
_COL_CANDIDATES_DATE = ['日期', '时间', 'date', 'time', '任务日期', 'task_date', '开始日期', '执行日期']
_COL_CANDIDATES_DESC = ['备注', '说明', '备注说明', 'desc', 'description', 'note', 'remark', 'notes', '详情']


def _smart_detect_columns(headers: List[str]):
    """Detect column indices for title / date / description by header name."""
    idx_title = None
    idx_date = None
    idx_desc = None
    for i, h in enumerate(headers):
        hl = str(h).strip().lower()
        if idx_title is None and any(k in hl for k in _COL_CANDIDATES_TITLE):
            idx_title = i
        if idx_date is None and any(k in hl for k in _COL_CANDIDATES_DATE):
            idx_date = i
        if idx_desc is None and any(k in hl for k in _COL_CANDIDATES_DESC):
            idx_desc = i
    if idx_title is None:
        # fallback: first column as title
        idx_title = 0
    return idx_title, idx_date, idx_desc


@router.post('/ai-import-preview')
async def ai_import_preview(
    file: UploadFile = File(...),
    _: User = Depends(require_admin),
):
    filename = file.filename or ""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    today = date.today()
    preview: List[dict] = []

    if ext in ('xlsx', 'xls'):
        content = await file.read()
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return {'preview': []}

        headers = [str(c) if c is not None else f'Column{i+1}' for i, c in enumerate(rows[0])]
        idx_title, idx_date, idx_desc = _smart_detect_columns(headers)

        current_trip_filter = ''
        for row in rows[1:]:
            vals = [str(c).strip() if c is not None else '' for c in row]
            title = vals[idx_title] if idx_title < len(vals) else ''
            if not title:
                continue

            # Detect group-header row: only title column has content → use as trip_filter
            non_empty_count = sum(1 for v in vals if v)
            if non_empty_count == 1:
                current_trip_filter = title
                continue

            task_d = today
            if idx_date is not None and idx_date < len(vals) and vals[idx_date]:
                try:
                    task_d = datetime.strptime(vals[idx_date], '%Y-%m-%d').date()
                except ValueError:
                    try:
                        task_d = datetime.strptime(vals[idx_date], '%Y/%m/%d').date()
                    except ValueError:
                        pass
            desc = vals[idx_desc] if idx_desc is not None and idx_desc < len(vals) else ''
            preview.append({
                'title': title,
                'task_date': task_d.isoformat(),
                'end_date': (task_d + timedelta(days=365)).isoformat(),
                'description': desc,
                'trip_filter': current_trip_filter,
            })

    elif ext in ('docx', 'doc'):
        content = await file.read()
        import docx
        doc = docx.Document(io.BytesIO(content))

        current_trip_filter = ''

        # Extract from paragraphs — detect bold headings as trip_filter groups
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect if paragraph is bold → use as group title
            is_bold = False
            if para.runs:
                bold_runs = [r for r in para.runs if r.text.strip()]
                if bold_runs:
                    is_bold = all(r.bold for r in bold_runs)

            if is_bold:
                current_trip_filter = text
                continue

            preview.append({
                'title': text,
                'task_date': today.isoformat(),
                'end_date': (today + timedelta(days=365)).isoformat(),
                'description': '',
                'trip_filter': current_trip_filter,
            })

        # Extract from tables — inherit last trip_filter
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                title = next((c for c in cells if c), '')
                if title:
                    preview.append({
                        'title': title,
                        'task_date': today.isoformat(),
                        'end_date': (today + timedelta(days=365)).isoformat(),
                        'description': '',
                        'trip_filter': current_trip_filter,
                    })

    else:
        raise HTTPException(status_code=400, detail='仅支持 .xlsx / .xls / .docx / .doc 文件')

    return {'preview': preview}


# ── AI 识别导入确认 ──

class AiImportConfirmBody(BaseModel):
    tasks: List[dict]


@router.post('/ai-import-confirm')
def ai_import_confirm(
    body: AiImportConfirmBody,
    db: Session = Depends(get_db),
    _: User = Depends(require_admin),
):
    created = []
    for item in body.tasks:
        title = item.get('title', '').strip()
        if not title:
            continue
        task_date_str = item.get('task_date', date.today().isoformat())
        try:
            td = date.fromisoformat(task_date_str)
            task_dt = datetime.combine(td, datetime.min.time())
        except (ValueError, TypeError):
            task_dt = datetime.combine(date.today(), datetime.min.time())

        end_date_str = item.get('end_date')
        end_dt = None
        if end_date_str:
            try:
                ed = date.fromisoformat(end_date_str)
                end_dt = datetime.combine(ed, datetime.min.time())
            except (ValueError, TypeError):
                pass
        if not end_dt:
            td_only = task_dt.date()
            end_dt = datetime.combine(td_only + timedelta(days=365), datetime.min.time())

        task = DailyTask(
            title=title,
            task_date=task_dt,
            end_date=end_dt,
            description=item.get('description', '') or '',
            trip_filter=item.get('trip_filter', '全部') or '全部',
            task_time=item.get('task_time', '') or '',
            end_time=item.get('end_time', '') or '',
            location=item.get('location', '') or '',
        )
        db.add(task)
        created.append(task)

    db.commit()
    return {'ok': True, 'count': len(created)}
